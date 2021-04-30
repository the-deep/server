import docx
import requests
import io
import re
import tempfile
import base64
import logging
from uuid import uuid4


from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml import OxmlElement, oxml_parser
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from PIL import Image

from utils.common import get_valid_xml_string


logger = logging.getLogger(__name__)


def xstr(text):
    return get_valid_xml_string(text, escape=False)


def _write_file(r, fp):
    for chunk in r.iter_content(chunk_size=1024):
        if chunk:
            fp.write(chunk)
    return fp


def _first_child_found_in(parent, tagnames):
    """
    Return the first child of parent with tag in *tagnames*, or None if
    not found.
    """
    for tagname in tagnames:
        child = parent.find(qn(tagname))
        if child is not None:
            return child
    return None


def _insert_element_before(parent, elm, successors):
    """
    Insert *elm* as child of *parent* before any existing child having
    tag name found in *successors*.
    """
    successor = _first_child_found_in(parent, successors)
    if successor is not None:
        successor.addprevious(elm)
    else:
        parent.append(elm)
    return elm


class Run:
    """
    Single run inside a paragraph
    """
    def __init__(self, ref):
        self.ref = ref

    def add_text(self, text):
        self.ref.add_text(xstr(text))
        return self

    def add_image(self, image):
        try:
            if image and len(image) > 0:
                fimage = tempfile.NamedTemporaryFile()
                if re.search(r'http[s]?://', image):
                    image = requests.get(image, stream=True, timeout=2)
                    _write_file(image, fimage)
                else:
                    image = base64.b64decode(image.split(',')[1])
                    fimage.write(image)
                self.ref.add_picture(fimage)
        except Exception:
            self.add_text('Invalid Image')

    def add_font_color(self, hex_color_string=None):
        hex_color_string = hex_color_string or '#000000'
        if '#' in hex_color_string:
            hex_color_string = hex_color_string[1:]
        color = RGBColor.from_string(hex_color_string)
        self.ref.font.color.rgb = color

    def add_shading(self, hex_color_string=None):
        """
        XML representation
            <w:shd w:fill="1F497D"/>
        """
        hex_color_string = hex_color_string or '#888888'
        if '#' in hex_color_string:
            hex_color_string = hex_color_string[1:]

        rPr = self.ref._element.get_or_add_rPr()
        ele = OxmlElement('w:shd')
        ele.set(qn('w:fill'), hex_color_string)
        rPr.append(ele)

    def add_inline_image(self, image, width, height):
        inline = self.ref.part.new_pic_inline(image, width, height)
        # Remove left/right spacing
        inline.set('distL', '0')
        inline.set('distR', '0')
        return self.ref._r.add_drawing(inline)

    def add_oval_shape(self, fill_hex_color=None):
        """
        https://python-docx.readthedocs.io/en/latest/user/shapes.html
        https://docs.microsoft.com/en-us/windows/win32/vml/web-workshop---specs---standards----how-to-use-vml-on-web-pages
        """
        fill_hex_color = fill_hex_color or '#ffffff'
        color = fill_hex_color
        if '#' != color[0]:
            color = '#' + color

        pict = OxmlElement('w:pict')
        nsmap = dict(
            v='urn:schemas-microsoft-com:vml'
        )
        oval_attrs = dict(
            id=str(uuid4()),
            style='width:12pt;height:12pt;z-index:-251658240;mso-position-vertical:top;mso-position-horizontal:left',
            fillcolor=color,
        )
        oval = oxml_parser.makeelement('{%s}%s' % (nsmap['v'], 'oval'),
                                       attrib=oval_attrs, nsmap=nsmap)

        border_attrs = dict(
            color='gray',
            joinstyle='round',
            endcap='flat'
        )
        stroke = oxml_parser.makeelement('{%s}%s' % (nsmap['v'], 'stroke'),
                                         attrib=border_attrs, nsmap=nsmap)
        oval.append(stroke)
        pict.append(oval)
        self.ref._element.append(pict)


class Paragraph:
    """
    One paragraph: supports normal text runs, hyperlinks,
    horizontal lines.
    """
    def __init__(self, ref):
        self.ref = ref

    def add_run(self, text=None, bold=False):
        run = Run(self.ref.add_run(xstr(text)))
        run.ref.bold = bold
        return run

    def add_hyperlink(self, url, text):
        part = self.ref.part
        r_id = part.relate_to(
            url,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True,
        )

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

        new_run = docx.oxml.shared.OxmlElement('w:r')
        r_pr = docx.oxml.shared.OxmlElement('w:rPr')

        new_run.append(r_pr)
        new_run.text = get_valid_xml_string(text)

        hyperlink.append(new_run)
        self.ref._p.append(hyperlink)

        r = docx.text.run.Run(new_run, self.ref)
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        return self

    def add_horizontal_line(self):
        p = self.ref._p
        p_pr = p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')

        _insert_element_before(p_pr, p_bdr, successors=(
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku',
            'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct',
            'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd',
            'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing',
            'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
        ))

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        p_bdr.append(bottom)

        return self

    def justify(self):
        self.ref.paragraph_format.alignment = \
            docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        return self

    def delete(self):
        p = self.ref._element
        p.getparent().remove(p)
        self.ref._p = self.ref._element = None

    def add_shaded_text(self, text, color):
        run = self.add_run(text)
        run.add_shading(color)

    def add_oval_shape(self, color):
        run = self.add_run(' ')
        run.add_oval_shape(color)

    def add_next_paragraph(self, other):
        p = self.ref._p
        p.addnext(other.ref._p)


class Document:
    """
    A docx document representation
    """
    def __init__(self, template=None):
        self.doc = docx.Document(template)

    def add_paragraph(self, text=None):
        return Paragraph(self.doc.add_paragraph(xstr(text)))

    def add_image(self, image):
        try:
            sec = self.doc.sections[-1]
            try:
                cols = int(
                    sec._sectPr.xpath('./w:cols')[0].get(qn('w:num'))
                )
                width = (
                    (sec.page_width / cols) -
                    (sec.right_margin + sec.left_margin)
                )
            except Exception:
                width = (
                    sec.page_width - (sec.right_margin + sec.left_margin)
                )

            if hasattr(image, 'read'):
                fimage = image
            elif image and len(image):
                fimage = tempfile.NamedTemporaryFile()
                if re.search(r'http[s]?://', image):
                    image = requests.get(image, stream=True, timeout=2)
                    _write_file(image, fimage)
                else:
                    image = base64.b64decode(image.split(',')[1])
                    fimage.write(image)

            image_width, _ = Image.open(fimage).size
            image_width = Pt(image_width)

            if image_width < width:
                self.doc.add_picture(fimage)
            else:
                self.doc.add_picture(fimage, width=width)
            self.doc.paragraphs[-1].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            return self
        except Exception:
            self.doc.add_paragraph('Invalid Image')
            logger.error(
                'export.formats.docx Add Image Error!!',
                exc_info=True,
            )
            return self

    def add_heading(self, text, level):
        return self.doc.add_heading(text, level=level)

    def add_page_break(self):
        self.doc.add_page_break()
        return self

    def save_to_file(self, fp):
        self.doc.save(fp)

    def save(self):
        buffer = io.BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()
