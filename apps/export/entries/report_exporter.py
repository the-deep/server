import os
import tempfile
import logging
from datetime import datetime
from subprocess import call

from django.conf import settings
from django.core.files.base import ContentFile, File
from django.db.models import (
    Case,
    When,
    Q,
)
from docx.shared import Inches
from deep.permalinks import Permalink
from utils.common import deep_date_parse, deep_date_format

from export.formats.docx import Document

from analysis_framework.models import Widget
from entry.models import (
    Entry,
    ExportData,
    Attribute,
    # EntryGroupLabel,
)
from entry.widgets import (
    scale_widget,
    time_widget,
    date_widget,
    time_range_widget,
    date_range_widget,
    geo_widget,
    select_widget,
    multiselect_widget,
    organigram_widget,
)
from entry.widgets.store import widget_store

from ary.export.affected_groups_info import get_affected_groups_info as ary_get_affected_groups_info
from ary.export.data_collection_techniques_info import (
    get_data_collection_techniques_info as ary_get_data_collection_techniques_info
)

from lead.models import Lead
from tabular.viz import renderer as viz_renderer
from export.models import Export

logger = logging.getLogger(__name__)

SEPARATOR = ', '
INTERNAL_SEPARATOR = '; '
ASSESSMENT_ICON_IMAGE_PATH = os.path.join(settings.BASE_DIR, 'apps/static/image/drop-icon.png')


class ExportDataVersionMismatch(Exception):
    pass


class WidgetExporter:

    @staticmethod
    def _add_common(para, text, bold):
        para.add_run(text, bold)

    @staticmethod
    def _add_scale_widget_data(para, label, color, bold):
        """
            Output: <oval shape with color> <space> <text>
        """
        para.add_oval_shape(color)
        para.add_run(label, bold)

    @classmethod
    def _get_scale_widget_data(cls, data, bold, **kwargs):
        """
        report for scale widget expects following keys
            - title
            - label
            - color
        as described here: apps.entry.widgets.scale_widget._get_scale
        """
        label = data.get('label')
        color = data.get('color')
        if label and color:
            return cls._add_scale_widget_data, label, color, bold

    @classmethod
    def _get_date_range_widget_data(cls, data, bold, **kwargs):
        """
        report for date range widget expects following
            - tuple (from, to)
        as described here: apps.entry.widgets.date_range_widget._get_date
        """
        date_renderer = kwargs['date_renderer']
        values = data.get('values', [])
        if len(values) == 2 and any(values):
            label = '{} - {}'.format(
                date_renderer(
                    deep_date_parse(
                        values[0],
                        raise_exception=False
                    ),
                    fallback="N/A",
                ),
                date_renderer(
                    deep_date_parse(
                        values[1],
                        raise_exception=False
                    ),
                    fallback="N/A",
                ),
            )
            return cls._add_common, label, bold

    @classmethod
    def _get_time_range_widget_data(cls, data, bold, **kwargs):
        """
        report for time range widget expects following
            - tuple (from, to)
        as described here: apps.entry.widgets.time_range_widget._get_time
        """
        values = data.get('values', [])
        if len(values) == 2 and any(values):
            text = '{} - {}'.format(
                values[0] or "~~:~~",
                values[1] or "~~:~~",
            )
            return cls._add_common, text, bold

    @classmethod
    def _get_date_widget_data(cls, data, bold, **kwargs):
        """
        report for date widget expects following
            - string (=date)
        as described here: apps.entry.widgets.date_widget
        """
        value = data.get('value')
        if not value:
            return
        date_renderer = kwargs['date_renderer']
        _value = date_renderer(deep_date_parse(value, raise_exception=False))
        if _value:
            return cls._add_common, _value, bold

    @classmethod
    def _get_time_widget_data(cls, data, bold, **kwargs):
        value = data.get('value')
        if value:
            return cls._add_common, value, bold

    @classmethod
    def _get_select_widget_data(cls, data, bold, **kwargs):
        type_ = data.get('type')
        value = [str(v) for v in data.get('value') or []]
        if type_ == 'list' and value:
            return cls._add_common, INTERNAL_SEPARATOR.join(value), bold

    @classmethod
    def _get_multi_select_widget_data(cls, data, bold, **kwargs):
        return cls._get_select_widget_data(data, bold, **kwargs)

    @classmethod
    def _get_organigram_widget_data(cls, data, bold, **kwargs):
        text = INTERNAL_SEPARATOR.join(
            '/'.join(value_with_parent)
            for value_with_parent in data.get('values')
        )
        return cls._add_common, text, bold

    @classmethod
    def _get_geo_widget_data(cls, data, bold, **kwargs):
        # XXX: Cache this value.
        # Right now everything needs to be loaded so doing this at entry save can take lot of memory
        geo_id_values = [str(v) for v in data.get('values') or []]
        if len(geo_id_values) == 0:
            return
        geo_values = kwargs['_get_geo_admin_level_1_data'](geo_id_values)
        if geo_values:
            return cls._add_common, geo_values, bold

    @classmethod
    def get_widget_information_into_report(
        cls,
        report,
        bold=True,
        **kwargs,
    ):
        """
        based on widget annotate information into report

        :param report: dict

        returns: function_to_add_to_report, *data
        """
        if not isinstance(report, dict):
            return
        if 'widget_id' in report:
            widget_id = report.get('widget_id')
            mapper = {
                scale_widget.WIDGET_ID: cls._get_scale_widget_data,
                date_range_widget.WIDGET_ID: cls._get_date_range_widget_data,
                time_range_widget.WIDGET_ID: cls._get_time_range_widget_data,
                time_widget.WIDGET_ID: cls._get_time_widget_data,
                date_widget.WIDGET_ID: cls._get_date_widget_data,
                geo_widget.WIDGET_ID: cls._get_geo_widget_data,
                select_widget.WIDGET_ID: cls._get_select_widget_data,
                organigram_widget.WIDGET_ID: cls._get_organigram_widget_data,
                multiselect_widget.WIDGET_ID: cls._get_multi_select_widget_data,
            }
            if widget_id in mapper.keys():
                if report.get('version') != widget_store[widget_id].DATA_VERSION:
                    raise ExportDataVersionMismatch(
                        f'{widget_id} widget data is not upto date. Export data being exported: {report}'
                    )
                return mapper[widget_id](report, bold, **kwargs)


class ReportExporter:
    def __init__(
        self,
        date_format: Export.DateFormat,
        citation_style: Export.CitationStyle,
        exporting_widgets=None,  # eg: ["517", "43", "42"]
        is_preview=False,
        show_lead_entry_id=True,
        show_assessment_data=True,
        show_entry_widget_data=True,
    ):
        self.is_preview = is_preview
        self.show_lead_entry_id = show_lead_entry_id
        self.show_assessment_data = show_assessment_data
        self.show_entry_widget_data = show_entry_widget_data
        # self.entry_group_labels = {}  # TODO: Remove entry group labels?

        self.doc = Document(
            os.path.join(settings.APPS_DIR, 'static/doc_export/template.docx')
        )
        self.lead_ids = []
        # ordered list of widget ids
        self.exporting_widgets_ids = [
            int(_id) for _id in exporting_widgets or []
        ]
        self.exporting_widgets_keys = list(
            Widget.objects.filter(id__in=self.exporting_widgets_ids).values_list('key', flat=True)
        )
        self.region_data = {}
        # XXX: Limit memory usage? (Or use redis?)
        self.geoarea_data_cache = {}
        self.assessment_data_cache = {}
        self.entry_widget_data_cache = {}

        # Citation
        self.citation_style = citation_style or Export.CitationStyle.DEFAULT

        # Date Format
        self.date_renderer = Export.get_date_renderer(date_format)

    def load_exportables(self, exportables, regions):
        exportables = exportables.filter(
            data__report__levels__isnull=False,
        )

        self.exportables = exportables

        geo_data_required = Widget.objects.filter(
            id__in=self.exporting_widgets_ids,
            widget_id=geo_widget.WIDGET_ID
        ).exists()
        # Load geo data if required
        if geo_data_required:
            self.region_data = {}
            for region in regions:
                # Collect geo area names for each admin level
                self.region_data[region.id] = [
                    {
                        'id': admin_level.id,
                        'level': admin_level.level,
                        'geo_area_titles': admin_level.get_geo_area_titles(),
                    }
                    for admin_level in region.adminlevel_set.all()
                ]

        return self

    def load_levels(self, levels):
        self.levels = levels
        return self

    def load_structure(self, structure):
        self.structure = structure
        return self

    def load_group_lables(self, entries, show_groups=False):
        # NOTE: This is causing high DB cpu usages. Removing as the feature doesn't exists in the UI.
        # FIXME: Analysize and fix the query.
        return self
        # if not show_groups:
        #     return self
        # for entry, group, label in EntryGroupLabel.objects.filter(entry__in=entries).order_by().values_list(
        #         'entry_id', 'group__title', 'label__title'):
        #     entry_d = self.entry_group_labels[entry] = self.entry_group_labels.get(entry, [])
        #     entry_d.append([group, label])
        # return self

    def load_text_from_text_widgets(self, entries, text_widget_ids):
        """
        Prefetch entry texts from text_Widgets
        text_widget_ids: ["id1", "id2"]
        """
        # User defined widget order (Generate order map)
        widget_ids = [int(id) for id in text_widget_ids]
        widget_map = {
            int(id): index
            for index, id in enumerate(text_widget_ids)
        }

        attribute_qs = Attribute.objects.filter(
            entry__in=entries,
            data__value__isnull=False,
            widget__id__in=widget_ids,
        )

        # Collect text from textWidgets for given entries using user defined widget data
        collected_widget_text = {}
        for (
            entry_id,
            text,
            widget_title,
            widget_id,
            widget_type,
        ) in attribute_qs.values_list(
            'entry_id',
            'data__value',  # Text
            'widget__title',
            'widget__id',
            'widget__widget_id',  # Widget Type
        ):
            if widget_type == 'conditionalWidget':
                continue

            widget_order = widget_map[widget_id]
            if widget_order is None:
                continue

            # Map text to entry->order->text
            collected_widget_text[entry_id] = collected_widget_text.get(entry_id) or {}
            collected_widget_text[entry_id][widget_order] = (
                widget_title,
                text and text.strip(),  # Widget Text Value
            )

        self.collected_widget_text = collected_widget_text
        return self

    def _generate_legend_page(self, project):
        para = self.doc.add_paragraph()
        self.legend_paragraph.add_next_paragraph(para)

        # todo in a table
        scale_widgets = project.analysis_framework.widget_set.filter(widget_id='scaleWidget')
        for widget in scale_widgets[::-1]:
            if not hasattr(widget, 'title'):
                continue
            title_para = self.doc.add_paragraph()
            title_para.ref.paragraph_format.right_indent = Inches(0.25)
            title_para.add_run(f'{widget.title}')
            for legend in widget.properties.get('options', [])[::-1]:
                para = self.doc.add_paragraph()
                para.ref.paragraph_format.right_indent = Inches(0.25)
                para.add_oval_shape(legend.get('color'))
                para.add_run(f'    {legend.get("label", "-Missing-")}')
                self.legend_paragraph.add_next_paragraph(para)
            self.legend_paragraph.add_next_paragraph(title_para)
        cond_widgets = project.analysis_framework.widget_set.filter(widget_id='conditionalWidget')
        for c_widget in cond_widgets[::-1]:
            for widget in filter(
                lambda x: x.get('widget', {}).get('widget_id') == 'scaleWidget',
                c_widget.properties.get('data', {}).get('widgets', [])
            ):
                if not widget.get('widget', {}).get('title'):
                    continue
                title_para = self.doc.add_paragraph()
                title_para.ref.paragraph_format.right_indent = Inches(0.25)
                title_para.add_run(f'{widget.get("widget", {}).get("title")}')
                for legend in widget.get('widget', {}).get('properties', {}).get('options', [])[::-1]:
                    para = self.doc.add_paragraph()
                    para.ref.paragraph_format.right_indent = Inches(0.25)
                    para.add_oval_shape(legend.get('color'))
                    para.add_run(f'    {legend.get("label", "-Missing-")}')
                    self.legend_paragraph.add_next_paragraph(para)
                self.legend_paragraph.add_next_paragraph(title_para)

    def _get_geo_admin_level_1_data(self, geo_id_values):
        if len(geo_id_values) == 0:
            return

        render_values = []
        for region_id, admin_levels in self.region_data.items():
            for admin_level in admin_levels:
                geo_area_titles = admin_level['geo_area_titles']
                for geo_id in geo_id_values:
                    if geo_id not in geo_area_titles:
                        continue
                    if geo_id in self.geoarea_data_cache:
                        title = self.geoarea_data_cache[geo_id]
                        title and render_values.append(title)
                        continue
                    self.geoarea_data_cache[geo_id] = None

                    title = geo_area_titles[geo_id].get('title')
                    parent_id = geo_area_titles[geo_id].get('parent_id')
                    if admin_level['level'] == 1:
                        title and render_values.append(title)
                        self.geoarea_data_cache[geo_id] = title
                        continue

                    # Try to look through parent
                    for _level in range(0, admin_level['level'])[::-1]:
                        if parent_id:
                            _geo_area_titles = admin_levels[_level]['geo_area_titles']
                            _geo_area = _geo_area_titles.get(parent_id) or {}
                            _title = _geo_area.get('title')
                            parent_id = _geo_area.get('parent_id')
                            if _level == 1:
                                _title and render_values.append(_title)
                                self.geoarea_data_cache[geo_id] = title
                                break

        if render_values:
            return INTERNAL_SEPARATOR.join(set(render_values))

    def _add_assessment_info_for_entry(self, assessment, para, bold=True):
        def _add_assessment_icon():
            # NOTE: Add icon here
            run = para.add_run('', bold=bold)
            with open(ASSESSMENT_ICON_IMAGE_PATH, 'rb') as fp:
                run.add_inline_image(fp, width=Inches(0.15), height=Inches(0.15))

        cache = self.assessment_data_cache.get(assessment.pk)

        # Check/Calculate data for assessment
        if cache is None:
            cache = {}
            # Collect Assessment GEO Data
            cache['locations'] = self._get_geo_admin_level_1_data(
                assessment.locations.values_list('id', flat=True),
            )
            cache['affected_groups_info'] = INTERNAL_SEPARATOR.join([
                '/'.join([str(s) for s in info.values() if s])
                for info in ary_get_affected_groups_info(assessment)['affected_groups_info']
            ])
            cache['data_collection_techniques_info'] = INTERNAL_SEPARATOR.join([
                f"{info['Sampling Size']} {info['Data Collection Technique']}"
                for info in ary_get_data_collection_techniques_info(assessment)['data_collection_technique']
                if info.get('Sampling Size')
            ])
            dc_start_date = deep_date_format(assessment.data_collection_start_date)
            dc_end_date = deep_date_format(assessment.data_collection_end_date)
            if dc_start_date or dc_end_date:
                cache['data_collection_date'] = f'Data collection: {dc_start_date} - {dc_end_date}'
            self.assessment_data_cache[assessment.pk] = cache

        locations = cache['locations']
        affected_groups_info = cache['affected_groups_info']
        data_collection_techniques_info = cache['data_collection_techniques_info']
        data_collection_date = cache['data_collection_date']

        to_process_fuctions = [
            func
            for condition, func in [
                (True, _add_assessment_icon),
                (locations, lambda: para.add_run(locations, bold=bold)),
                (affected_groups_info, lambda: para.add_run(affected_groups_info, bold=bold)),
                (data_collection_techniques_info, lambda: para.add_run(data_collection_techniques_info, bold=bold)),
                (data_collection_date, lambda: para.add_run(data_collection_date, bold=bold)),
            ] if condition
        ]

        para.add_run(' [', bold=True)
        # Finally add all assessment data to the docx
        total_process_functions = len(to_process_fuctions) - 1
        for index, add_data in enumerate(to_process_fuctions):
            add_data()
            if index < total_process_functions:
                para.add_run(SEPARATOR, bold=bold)
        para.add_run('] ', bold=True)

    def _generate_for_entry_widget_data(self, entry, para):
        if entry.id not in self.entry_widget_data_cache:
            raw_export_data = []
            for each in entry.exportdata_set.all():
                export_datum = {
                    **(each.data.get('common') or {}),
                    **(each.data.get('report') or {}),
                }
                if export_datum.get('widget_key') and export_datum['widget_key'] in self.exporting_widgets_keys:
                    raw_export_data.append(export_datum)
            raw_export_data.sort(key=lambda x: self.exporting_widgets_keys.index(x['widget_key']))

            export_data = []
            if raw_export_data:
                for data in raw_export_data:
                    try:
                        if resp := WidgetExporter.get_widget_information_into_report(
                            data,
                            bold=True,
                            _get_geo_admin_level_1_data=self._get_geo_admin_level_1_data,
                            date_renderer=self.date_renderer,
                        ):
                            export_data.append(resp)
                    except ExportDataVersionMismatch:
                        logger.error(
                            f'ExportDataVersionMismatch: For entry {entry.id}, project {entry.project.id}',
                            exc_info=True
                        )
            self.entry_widget_data_cache[entry.id] = export_data
        export_data = self.entry_widget_data_cache[entry.id]

        if export_data:
            para.add_run(' [', bold=True)
            export_data_len = len(export_data) - 1
            for index, [func, *args] in enumerate(export_data):
                func(para, *args)  # Add to para
                if index < export_data_len:
                    para.add_run(SEPARATOR, bold=True)
            para.add_run('] ', bold=True)

    def _generate_for_entry(self, entry):
        """
        Generate paragraphs for an entry
        """

        para = self.doc.add_paragraph().justify()

        # entry-lead id
        if self.show_lead_entry_id:
            para.add_run('[', bold=True)
            # Add lead-entry id
            url = Permalink.entry(entry.project_id, entry.lead_id, entry.id)
            para.add_hyperlink(url, f"{entry.lead_id}-{entry.id}")
            para.add_run(']', bold=True)

        # Assessment Data
        if self.show_assessment_data and getattr(entry.lead, 'assessment', None):
            self._add_assessment_info_for_entry(entry.lead.assessmentregistry, para, bold=True)

        # Entry widget Data
        if self.show_entry_widget_data:
            self._generate_for_entry_widget_data(entry, para)

        # Format is
        # excerpt (source) OR excerpt \n text from widgets \n (source)
        # where source is hyperlinked to appropriate url

        # Excerpt can also be image
        excerpt = (
            entry.excerpt if entry.entry_type == Entry.TagType.EXCERPT
            else ''
        )

        if self.citation_style == Export.CitationStyle.STYLE_1:
            para.add_run(excerpt.rstrip('.'))
        else:  # Default
            para.add_run(excerpt)

        # Add texts from TextWidget
        entry_texts = self.collected_widget_text.get(entry.id) or {}
        widget_texts_exists = not not entry_texts.keys()
        if widget_texts_exists:
            self.doc.add_paragraph()  # Blank line
            for order in sorted(entry_texts.keys()):
                title, text = entry_texts[order]
                para = self.doc.add_paragraph().justify()
                para.add_run(f'{title}: ', bold=True)
                para.add_run(text)
            para = self.doc.add_paragraph().justify()

        # NOTE: Use doc.add_image for limiting image size to page width and run.add_image for actual size image

        image = None
        image_text = None
        if entry.entry_type == Entry.TagType.IMAGE:
            image = (entry.image and entry.image.file) or entry.image_raw
            # para.add_run().add_image(entry.image_raw)
        elif entry.entry_type == Entry.TagType.DATA_SERIES and entry.tabular_field:
            image = viz_renderer.get_entry_image(entry)
            h_stats = (entry.tabular_field.cache or {}).get('health_stats', {})

            image_text = ' Total values: {}'.format(h_stats.get('total', 'N/A'))
            for key in ['invalid', 'null']:
                if h_stats.get(key):
                    image_text += f', {key.title()} values: {h_stats.get(key)}' if h_stats.get(key) else ''

        if entry.entry_type == Entry.TagType.ATTACHMENT:
            image = entry.entry_attachment.file_preview
        if image:
            self.doc.add_image(image)
            if image_text:
                self.doc.add_paragraph(image_text).justify()
            para = self.doc.add_paragraph().justify()

        lead = entry.lead
        self.lead_ids.append(lead.id)

        # --- Reference Start
        if widget_texts_exists or image:
            para.add_run('(')  # Starting from new line
        else:
            para.add_run(' (')  # Starting from same line

        if self.citation_style == Export.CitationStyle.STYLE_1:
            source = ''
            author = lead.get_authors_display(short_name=True)
        else:  # Default
            source = lead.get_source_display() or 'Reference'
            author = lead.get_authors_display()

        url = lead.url or Permalink.lead_share_view(lead.uuid)
        date = entry.lead.published_on

        if self.citation_style == Export.CitationStyle.STYLE_1:
            # Add author is available
            if (author and author.lower() != (source or '').lower()):
                if url:
                    para.add_hyperlink(url, f'{author} ')
                else:
                    para.add_run(f'{author} ')
        else:  # Default
            # Add author is available
            if (author and author.lower() != (source or '').lower()):
                if url:
                    para.add_hyperlink(url, f'{author}, ')
                else:
                    para.add_run(f'{author}, ')
            # Add source (with url if available)
            if url:
                para.add_hyperlink(url, source)
            else:
                para.add_run(source)

        # Add (confidential/restricted) to source without ,
        if lead.confidentiality == Lead.Confidentiality.CONFIDENTIAL:
            para.add_run(' (confidential)')
        elif lead.confidentiality == Lead.Confidentiality.RESTRICTED:
            para.add_run(' (restricted)')

        if self.citation_style == Export.CitationStyle.STYLE_1:
            pass
        else:  # Default
            # Add lead title if available
            if lead.title:
                para.add_run(f", {lead.title}")

        # Finally add date
        if date:
            if self.citation_style == Export.CitationStyle.STYLE_1:
                para.add_run(f" {self.date_renderer(date)}")
            else:  # Default
                para.add_run(f", {self.date_renderer(date)}")

        para.add_run(')')
        # --- Reference End

        if self.citation_style == Export.CitationStyle.STYLE_1:
            para.add_run('.')

        self.doc.add_paragraph()

    def _load_into_levels(
            self,
            entry,
            keys,
            levels,
            entries_map,
            valid_levels,
    ):
        """
        Map an entry into corresponding levels
        """
        parent_level_valid = False
        for level in levels:
            level_id = level.get('id')
            valid_level = (level_id in keys)

            if valid_level:
                if not entries_map.get(level_id):
                    entries_map[level_id] = []
                entries_map[level_id].append(entry)

            sublevels = level.get('sublevels')
            if sublevels:
                valid_level = valid_level or self._load_into_levels(
                    entry,
                    keys,
                    sublevels,
                    entries_map,
                    valid_levels,
                )

            if valid_level and level_id not in valid_levels:
                valid_levels.append(level_id)

            parent_level_valid = parent_level_valid or valid_level
        return parent_level_valid

    def _generate_for_levels(
            self,
            levels,
            level_entries_map,
            valid_levels,
            structures=None,
            heading_level=2,
    ):
        """
        Generate paragraphs for all entries in this level and recursively
        do it for further sublevels
        """
        if structures is not None:
            level_map = dict((level.get('id'), level) for level in levels)
            levels = [level_map[s['id']] for s in structures]

        for level in levels:
            if level.get('id') not in valid_levels:
                continue

            title = level.get('title')
            entries = level_entries_map.get(level.get('id'))
            sublevels = level.get('sublevels')

            if entries or sublevels:
                self.doc.add_heading(title, heading_level)
                self.doc.add_paragraph()

            if entries:
                iterable_entries = entries[:Export.PREVIEW_ENTRY_SIZE] if self.is_preview else entries
                [self._generate_for_entry(entry) for entry in iterable_entries]

            if sublevels:
                substructures = None
                if structures:
                    substructures = next((
                        s.get('levels') for s in structures
                        if s['id'] == level.get('id')
                    ), None)

                self._generate_for_levels(
                    sublevels,
                    level_entries_map,
                    valid_levels,
                    substructures,
                    heading_level + 1,
                )

    def _generate_for_uncategorized(self, entries, categorized_entry_processed):
        entries = entries.exclude(
            Q(exportdata__data__report__keys__isnull=False) |
            Q(exportdata__data__report__keys__len__gt=0)
        )

        if entries.count() == 0:
            return

        if self.is_preview and categorized_entry_processed >= Export.PREVIEW_ENTRY_SIZE:
            return

        self.doc.add_heading('Uncategorized', 2)
        self.doc.add_paragraph()

        iterable_entries = entries[:Export.PREVIEW_ENTRY_SIZE - categorized_entry_processed] if self.is_preview else entries
        for entry in iterable_entries:
            self._generate_for_entry(entry)

    def pre_build_document(self, project):
        """
        Structure the document
        """
        self.doc.add_heading(
            'DEEP Export — {} — {}'.format(
                self.date_renderer(datetime.today()),
                project.title,
            ),
            1,
        )
        self.doc.add_paragraph()

        self.legend_heading = self.doc.add_heading('Legends', 2)
        self.legend_paragraph = self.doc.add_paragraph()

    def add_entries(self, entries):
        """
        Add entries and generate parapgraphs for all entries
        """
        if entries:
            self.pre_build_document(entries[0].project)
        exportables = self.exportables
        af_levels_map = dict((str(level.get('id')), level.get('levels')) for level in self.levels)
        uncategorized = False
        categorized_entry_processed = 0  # NOTE: Used for preview limit only

        if self.structure:
            ids = [s['id'] for s in self.structure]
            uncategorized = 'uncategorized' in ids
            ids = [id for id in ids if id != 'uncategorized']

            order = Case(*[
                When(pk=pk, then=pos)
                for pos, pk
                in enumerate(ids)
            ])
            exportables = exportables.filter(pk__in=ids).order_by(order)

        for exportable in exportables:
            levels = (
                # Custom levels provided by client
                af_levels_map.get(str(exportable.pk)) or
                # Predefined levels available in server
                exportable.data.get('report').get('levels')
            )

            level_entries_map = {}
            valid_levels = []

            iterable_entries = entries[:Export.PREVIEW_ENTRY_SIZE] if self.is_preview else entries
            for entry in iterable_entries:
                # TODO
                # Set entry.report_data to all exportdata for all exportable
                # for this entry for later use

                export_data = ExportData.objects.filter(
                    entry=entry,
                    exportable=exportable,
                    data__report__keys__isnull=False,
                ).first()

                if export_data:
                    self._load_into_levels(
                        entry, export_data.data.get('report').get('keys'),
                        levels, level_entries_map, valid_levels,
                    )
                    categorized_entry_processed += 1

            structures = self.structure and next((
                s.get('levels') for s in self.structure
                if str(s['id']) == str(exportable.id)
            ), None)
            self._generate_for_levels(levels, level_entries_map,
                                      valid_levels, structures)

        if uncategorized:
            self._generate_for_uncategorized(entries, categorized_entry_processed)
        if entries:
            self._generate_legend_page(entries[0].project)

        return self

    def export(self, pdf=False):
        """
        Export and return export data
        """

        # Get all leads to generate Bibliography
        lead_ids = list(set(self.lead_ids))
        leads = Lead.objects.filter(id__in=lead_ids)

        self.doc.add_paragraph().add_horizontal_line()
        self.doc.add_paragraph()
        self.doc.add_heading('Bibliography', 1)
        self.doc.add_paragraph()

        for lead in leads:
            # Format is"
            # Source. Title. Date. Url

            para = self.doc.add_paragraph()
            author = lead.get_authors_display()
            source = lead.get_source_display() or 'Missing source'

            if author:
                para.add_run(f'{author}.')
            para.add_run(f' {source}.')
            para.add_run(f' {lead.title}.')
            if lead.published_on:
                para.add_run(f" {self.date_renderer(lead.published_on)}. ")

            para = self.doc.add_paragraph()
            url = lead.url or Permalink.lead_share_view(lead.uuid)
            if url:
                para.add_hyperlink(url, url)
            else:
                para.add_run('Missing url.')

            if lead.confidentiality == Lead.Confidentiality.CONFIDENTIAL:
                para.add_run(' (confidential)')
            elif lead.confidentiality == Lead.Confidentiality.RESTRICTED:
                para.add_run(' (restricted)')

            self.doc.add_paragraph()
        # self.doc.add_page_break()

        if pdf:
            temp_doc = tempfile.NamedTemporaryFile(dir=settings.TEMP_DIR)
            self.doc.save_to_file(temp_doc)

            filename = temp_doc.name.split('/')[-1]
            temp_pdf = os.path.join(settings.TEMP_DIR, '{}.pdf'.format(filename))

            call(['libreoffice', '--headless', '--convert-to', 'pdf', temp_doc.name, '--outdir', settings.TEMP_DIR])
            file = File(open(temp_pdf, 'rb'))

            # Cleanup
            os.unlink(temp_pdf)
            temp_doc.close()

        else:
            buffer = self.doc.save()
            file = ContentFile(buffer)

        return file
